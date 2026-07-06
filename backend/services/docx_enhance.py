from io import BytesIO
import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from services.barcode_gen import generate_barcode

LOGO_PLACEHOLDER = "{{company_logo}}"
REF_PLACEHOLDER = "{{ref_number}}"
LOGO_WIDTH_INCHES = 1.4  # fixed — do not change
BARCODE_WIDTH_INCHES = 1.65

def _replace_logo_in_paragraph(paragraph, logo_bytes: bytes) -> bool:
    full_text = "".join(run.text for run in paragraph.runs)
    if LOGO_PLACEHOLDER not in full_text:
        return False
    paragraph.clear()
    run = paragraph.add_run()
    run.add_picture(BytesIO(logo_bytes), width=Inches(LOGO_WIDTH_INCHES))
    return True


def _scan_paragraphs(paragraphs, logo_bytes: bytes) -> bool:
    changed = False
    for paragraph in paragraphs:
        if _replace_logo_in_paragraph(paragraph, logo_bytes):
            changed = True
    return changed


def inject_logo(docx_path: str, image_path: str) -> None:
    if not image_path or not os.path.exists(image_path):
        return

    with open(image_path, "rb") as f:
        logo_bytes = f.read()

    doc = Document(docx_path)
    changed = _scan_paragraphs(doc.paragraphs, logo_bytes)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if _scan_paragraphs(cell.paragraphs, logo_bytes):
                    changed = True

    for section in doc.sections:
        header = section.header
        if header:
            if _scan_paragraphs(header.paragraphs, logo_bytes):
                changed = True
            for table in header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if _scan_paragraphs(cell.paragraphs, logo_bytes):
                            changed = True

        for extra_header in (section.first_page_header, section.even_page_header):
            if extra_header:
                if _scan_paragraphs(extra_header.paragraphs, logo_bytes):
                    changed = True
                for table in extra_header.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if _scan_paragraphs(cell.paragraphs, logo_bytes):
                                changed = True

        for footer in (section.footer, section.first_page_footer, section.even_page_footer):
            if footer and _scan_paragraphs(footer.paragraphs, logo_bytes):
                changed = True

    if changed:
        doc.save(docx_path)


def _paragraph_has_ref_placeholder(paragraph) -> bool:
    return REF_PLACEHOLDER in "".join(run.text for run in paragraph.runs)


def _replace_ref_in_paragraph(paragraph, ref_number: str, barcode_bytes: bytes | None) -> bool:
    if not _paragraph_has_ref_placeholder(paragraph):
        return False

    paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.space_after = Pt(0)

    if barcode_bytes:
        run_img = paragraph.add_run()
        run_img.add_picture(BytesIO(barcode_bytes), width=Inches(BARCODE_WIDTH_INCHES))
    else:
        run_ref = paragraph.add_run()
        run_ref.text = ref_number
        run_ref.font.size = Pt(12)
        run_ref.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    return True


def _scan_body_for_ref(paragraphs, ref_number: str, barcode_bytes: bytes | None) -> bool:
    for paragraph in paragraphs:
        if _replace_ref_in_paragraph(paragraph, ref_number, barcode_bytes):
            return True
    return False


def inject_ref_barcode(docx_path: str, ref_number: str, issue_date: str | None = None) -> None:
    if not ref_number:
        return

    barcode_bytes = None
    try:
        barcode_bytes = generate_barcode(ref_number)
    except Exception:
        barcode_bytes = None

    doc = Document(docx_path)
    changed = _scan_body_for_ref(doc.paragraphs, ref_number, barcode_bytes)

    if not changed:
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if _scan_body_for_ref(cell.paragraphs, ref_number, barcode_bytes):
                        changed = True
                        break
                if changed:
                    break
            if changed:
                break

    if changed:
        doc.save(docx_path)

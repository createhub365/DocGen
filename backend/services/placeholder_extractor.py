import re

from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

from services.docx_xml_fill import extract_placeholder_ids_from_docx

PLACEHOLDER_PATTERN = re.compile(r"\{\{(\w+)\}\}")


def auto_label(placeholder_id: str) -> str:
    return placeholder_id.replace("_", " ").title()


def merge_paragraph_runs(paragraph):
    """Merge split runs so placeholder regex sees full {{id}} tokens."""
    if not hasattr(paragraph, "runs") or not paragraph.runs:
        return
    full_text = "".join(run.text for run in paragraph.runs)
    for i, run in enumerate(paragraph.runs):
        run.text = full_text if i == 0 else ""


def get_all_paragraphs(doc):
    """
    Collect ALL paragraphs from every possible location in a .docx:
    - Body paragraphs
    - All table cells (including nested tables)
    - All header types (default, first page, even page)
    - All footer types (default, first page, even page)
    - Text boxes and shapes (via XML fallback)
    """
    paragraphs = []

    paragraphs.extend(doc.paragraphs)

    def extract_table_paragraphs(table):
        for row in table.rows:
            for cell in row.cells:
                paragraphs.extend(cell.paragraphs)
                for nested_table in cell.tables:
                    extract_table_paragraphs(nested_table)

    for table in doc.tables:
        extract_table_paragraphs(table)

    for section in doc.sections:
        header_footer_objects = []

        if section.header and not section.header.is_linked_to_previous:
            header_footer_objects.append(section.header)
        if section.footer and not section.footer.is_linked_to_previous:
            header_footer_objects.append(section.footer)

        try:
            if section.first_page_header:
                header_footer_objects.append(section.first_page_header)
            if section.first_page_footer:
                header_footer_objects.append(section.first_page_footer)
        except Exception:
            pass

        try:
            if section.even_page_header:
                header_footer_objects.append(section.even_page_header)
            if section.even_page_footer:
                header_footer_objects.append(section.even_page_footer)
        except Exception:
            pass

        for hf in header_footer_objects:
            paragraphs.extend(hf.paragraphs)
            for table in hf.tables:
                extract_table_paragraphs(table)

    try:
        body_xml = doc.element.body
        for txbx in body_xml.iter(qn("w:txbxContent")):
            for p_elem in txbx.iter(qn("w:p")):
                try:
                    paragraphs.append(Paragraph(p_elem, doc))
                except Exception:
                    texts = []
                    for t in p_elem.iter(qn("w:t")):
                        if t.text:
                            texts.append(t.text)
                    raw_text = "".join(texts)
                    if "{{" in raw_text:

                        class _RawPara:
                            def __init__(self, text):
                                self.text = text
                                self.runs = []

                        paragraphs.append(_RawPara(raw_text))
    except Exception:
        pass

    return paragraphs


def _extract_ids_from_document(docx_path: str) -> list[str]:
    doc = Document(docx_path)
    seen: set[str] = set()
    ordered: list[str] = []

    for paragraph in get_all_paragraphs(doc):
        merge_paragraph_runs(paragraph)
        text = getattr(paragraph, "text", "") or "".join(
            run.text for run in getattr(paragraph, "runs", [])
        )
        for match in PLACEHOLDER_PATTERN.finditer(text):
            pid = match.group(1)
            if pid not in seen:
                seen.add(pid)
                ordered.append(pid)

    return ordered


def extract_placeholders(docx_path: str, label_overrides: dict = None) -> list:
    """Read placeholders from Word document (XML + python-docx scan)."""
    label_overrides = label_overrides or {}

    xml_ids = extract_placeholder_ids_from_docx(docx_path)
    doc_ids = _extract_ids_from_document(docx_path)

    seen: set[str] = set()
    ordered_ids: list[str] = []
    for pid in xml_ids + doc_ids:
        if pid not in seen:
            seen.add(pid)
            ordered_ids.append(pid)

    return [
        {
            "id": pid,
            "label": label_overrides.get(pid) or auto_label(pid),
            "type": "text",
            "required": True,
        }
        for pid in ordered_ids
    ]


if __name__ == "__main__":
    import sys

    if len(sys.argv) > 1:
        placeholders = extract_placeholders(sys.argv[1])
        print(f"Found {len(placeholders)} placeholders:")
        for p in placeholders:
            print(f"  {p['id']:30} → {p['label']}")

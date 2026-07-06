import json
import os
import sys

from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from database import SessionLocal, engine, Base
import models
from auth import hash_password

load_dotenv()
TEMPLATE_DIR = os.getenv("TEMPLATE_DIR", "./template_store")

SEED_DATA = {
    "ca": {
        "name": "Canada",
        "trades": {
            "Construction Worker": ["BuildRight Corp", "NorthBuild Ltd", "CanCon Inc"],
            "Truck Driver": ["TransCan Logistics", "MapleFreight", "NorthStar Transport"],
            "Welder": ["SteelForce Canada", "WeldPro Ltd", "NorthMetal Co"],
            "Electrician": ["PowerGrid Canada", "ElecPro Inc", "CanElect Ltd"],
            "Plumber": ["FlowRight Canada", "NorthPipe Co", "AquaPro Ltd"],
            "Nurse": ["Toronto General Hospital", "BC Health Authority", "Alberta Care"],
            "Cook": ["CanCuisine Corp", "FoodPro Ltd", "NorthKitchen Inc"],
        },
    },
    "au": {
        "name": "Australia",
        "trades": {
            "Construction Worker": ["AusBuild Pty", "SydneyConstruct", "OzBuilders Ltd"],
            "Welder": ["AusWeld Pty", "MetalForce AU", "SteelOz Ltd"],
            "Electrician": ["PowerAus Pty", "ElecOz Ltd", "VoltPro AU"],
            "IT Engineer": ["TechOz Pty", "AusIT Solutions", "Digital Futures AU"],
            "Cook": ["OzCuisine Pty", "AusFood Ltd", "SydneyKitchen Co"],
            "Nurse": ["Sydney Hospital Group", "Melbourne Health", "AusCare Ltd"],
        },
    },
    "uk": {
        "name": "United Kingdom",
        "trades": {
            "Construction Worker": ["UK Build Ltd", "LondonConstruct", "BritBuild Co"],
            "Plumber": ["FlowUK Ltd", "BritPipe Co", "LondonPlumb"],
            "Nurse": ["NHS Partners Ltd", "London Medical Group", "BritCare Ltd"],
            "IT Engineer": ["TechUK Ltd", "LondonTech", "BritSoft Co"],
            "Accountant": ["FinanceUK Ltd", "LondonAccounts", "BritFinance Co"],
        },
    },
    "uae": {
        "name": "UAE",
        "trades": {
            "Construction Worker": ["Emirates Build LLC", "Dubai Construct", "AbuDhabi Build Co"],
            "Engineer": ["Emirates Engineering LLC", "Gulf Tech Co", "Dubai Solutions"],
            "Driver": ["Emirates Transport", "Gulf Logistics", "Dubai Drive LLC"],
            "Security Guard": ["Gulf Security LLC", "Emirates Guard", "Dubai Protect"],
            "Cook": ["Emirates Cuisine LLC", "Gulf Food Co", "Dubai Kitchen LLC"],
            "Cleaner": ["Emirates Facilities LLC", "Gulf Housekeeping", "Dubai Clean Co"],
        },
    },
    "nz": {
        "name": "New Zealand",
        "trades": {
            "Construction Worker": ["KiwiBuild Ltd", "NZConstruct", "SouthBuild Co"],
            "Carpenter": ["KiwiCraft Ltd", "NZWood Co", "SouthCarpentry"],
            "Electrician": ["KiwiPower Ltd", "NZElec Co", "SouthVolt"],
            "Nurse": ["Auckland Health Group", "Wellington Medical", "NZCare Ltd"],
            "Cook": ["KiwiCuisine Ltd", "NZFood Co", "SouthKitchen"],
        },
    },
}

DOCUMENT_TYPES = [
    ("Offer Letter", "offer_letter"),
    ("Demand Letter", "demand_letter"),
    ("Employment Contract", "employment_contract"),
    ("Appointment Letter", "appointment_letter"),
]

LABEL_OVERRIDES = {
    "cand_name": "Candidate Full Name",
    "cand_dob": "Date of Birth",
    "cand_passport": "Passport Number",
    "joining_date": "Joining Date",
    "trade": "Trade Category",
}


def create_employment_contract_docx(path: str):
    doc = Document()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("{{company_name}}")
    run.bold = True
    run.font.size = Pt(18)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Employment Contract")
    run.font.size = Pt(14)

    doc.add_paragraph()
    doc.add_paragraph("Reference: {{ref_number}}")
    doc.add_paragraph("Date: {{offer_date}}")
    doc.add_paragraph()

    doc.add_paragraph("This Employment Contract is entered into between {{company_name}} and {{cand_name}}.")
    doc.add_paragraph()

    doc.add_paragraph("Employee Details:")
    doc.add_paragraph("Full Name: {{cand_name}}")
    doc.add_paragraph("Date of Birth: {{cand_dob}}")
    doc.add_paragraph("Passport Number: {{cand_passport}}")
    doc.add_paragraph("Nationality: {{cand_nationality}}")
    doc.add_paragraph()

    doc.add_paragraph("Position: {{position}}")
    doc.add_paragraph("Department: {{department}}")
    doc.add_paragraph("Work Location: {{location}}")
    doc.add_paragraph("Joining Date: {{joining_date}}")
    doc.add_paragraph("Contract Duration: {{duration}}")
    doc.add_paragraph("Working Hours: {{working_hours}}")
    doc.add_paragraph("Probation Period: {{probation}}")
    doc.add_paragraph("Salary: {{salary}}")
    doc.add_paragraph()

    doc.add_paragraph("Employer Address: {{company_address}}")
    doc.add_paragraph()
    doc.add_paragraph("Signed by authorized representative of {{company_name}}.")

    os.makedirs(os.path.dirname(path) or ".", exist_ok=True)
    doc.save(path)


EMPLOYMENT_CONTRACT_LABELS = {
    "cand_name": "Candidate Full Name",
    "cand_passport": "Passport Number",
    "cand_dob": "Date of Birth",
    "cand_nationality": "Nationality",
    "position": "Position Title",
    "department": "Department",
    "salary": "Salary",
    "joining_date": "Joining Date",
    "duration": "Contract Duration",
    "working_hours": "Working Hours",
    "probation": "Probation Period",
    "location": "Work Location",
    "company_name": "Company Name",
    "company_address": "Company Address",
    "ref_number": "Reference Number",
    "offer_date": "Contract Date",
}


APPOINTMENT_LETTER_LABELS = {
    "cand_name": "Candidate Full Name",
    "reporting_to": "Reporting Manager",
    "joining_date": "Date of Joining",
    "offer_date": "Letter Date",
    "probation_period": "Probation Duration",
    "notice_period": "Notice Period",
}


def create_appointment_letter_docx(path: str):
    doc = Document()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("{{company_name}}")
    run.bold = True
    run.font.size = Pt(18)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Appointment Letter")
    run.font.size = Pt(14)

    doc.add_paragraph()
    doc.add_paragraph("Reference: {{ref_number}}")
    doc.add_paragraph("Date: {{offer_date}}")
    doc.add_paragraph()

    doc.add_paragraph("Dear {{cand_name}},")
    doc.add_paragraph()

    doc.add_paragraph(
        "We are pleased to confirm your appointment as {{position}} in the "
        "{{department}} department at {{company_name}}."
    )
    doc.add_paragraph()

    doc.add_paragraph("Position: {{position}}")
    doc.add_paragraph("Department: {{department}}")
    doc.add_paragraph("Work Location: {{work_location}}")
    doc.add_paragraph("Reporting To: {{reporting_to}}")
    doc.add_paragraph("Date of Joining: {{joining_date}}")
    doc.add_paragraph("Salary: {{salary}}")
    doc.add_paragraph("Probation Period: {{probation_period}}")
    doc.add_paragraph("Notice Period: {{notice_period}}")
    doc.add_paragraph()

    doc.add_paragraph(
        "Please report to {{reporting_to}} on {{joining_date}} at {{work_location}}."
    )
    doc.add_paragraph()
    doc.add_paragraph("Sincerely,")
    doc.add_paragraph()
    doc.add_paragraph("_________________________")
    doc.add_paragraph("Human Resources")
    doc.add_paragraph("{{company_name}}")

    os.makedirs(os.path.dirname(path) or ".", exist_ok=True)
    doc.save(path)


def ensure_appointment_letter(db):
    """Add Appointment Letter doc type + sample template if missing (existing DBs)."""
    appt_type = (
        db.query(models.DocumentType)
        .filter(models.DocumentType.slug == "appointment_letter")
        .first()
    )
    if not appt_type:
        appt_type = models.DocumentType(name="Appointment Letter", slug="appointment_letter")
        db.add(appt_type)
        db.flush()

    existing_template = (
        db.query(models.Template)
        .filter(models.Template.document_type_id == appt_type.id)
        .first()
    )
    if existing_template:
        return

    canada = db.query(models.Country).filter(models.Country.code == "ca").first()
    if not canada:
        return

    construction = (
        db.query(models.Trade)
        .filter(models.Trade.country_id == canada.id, models.Trade.name == "Construction Worker")
        .first()
    )
    buildright = (
        db.query(models.Company)
        .filter(
            models.Company.country_id == canada.id,
            models.Company.trade_id == construction.id if construction else None,
            models.Company.name == "BuildRight Corp",
        )
        .first()
    )
    if not construction or not buildright:
        return

    appt_filename = "sample_canada_construction_buildrightcorp_appointment_letter.docx"
    appt_path = os.path.join(TEMPLATE_DIR, appt_filename)
    create_appointment_letter_docx(appt_path)

    template = models.Template(
        document_type_id=appt_type.id,
        company_id=buildright.id,
        trade_id=construction.id,
        country_id=canada.id,
        docx_filename=appt_filename,
        label_overrides_json=json.dumps(APPOINTMENT_LETTER_LABELS),
        category="Construction and Infrastructure",
        format_slug="format1",
        format_label="Format 1 — Corporate Formal (Navy & Gold)",
        version=1,
        is_active=True,
    )
    db.add(template)
    db.commit()
    print("Appointment Letter template seeded.")


def create_sample_docx(path: str):
    doc = Document()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("BuildRight Corp")
    run.bold = True
    run.font.size = Pt(18)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Official Offer of Employment")
    run.font.size = Pt(14)

    doc.add_paragraph()
    date_para = doc.add_paragraph("Date: {{joining_date}}")
    doc.add_paragraph()

    doc.add_paragraph("Dear {{cand_name}},")
    doc.add_paragraph()

    body_text = (
        "We are pleased to offer you the position of {{position}} with BuildRight Corp. "
        "This offer is made in connection with your application under the {{trade}} category "
        "for employment in {{location}}."
    )
    doc.add_paragraph(body_text)
    doc.add_paragraph()

    doc.add_paragraph(
        "Your anticipated start date is {{joining_date}}. The contract duration will be "
        "{{duration}}, with a compensation package of {{salary}}."
    )
    doc.add_paragraph()

    doc.add_paragraph(
        "For immigration processing purposes, please note your passport number "
        "{{cand_passport}} and date of birth {{cand_dob}} as registered in our records."
    )
    doc.add_paragraph()

    doc.add_paragraph(
        "Please sign and return a copy of this letter to confirm your acceptance of this offer."
    )
    doc.add_paragraph()
    doc.add_paragraph("Sincerely,")
    doc.add_paragraph()
    doc.add_paragraph("_________________________")
    doc.add_paragraph("Human Resources Manager")
    doc.add_paragraph("BuildRight Corp")

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run(
        "BuildRight Corp | 123 Construction Ave, Toronto, ON | www.buildrightcorp.ca"
    )

    os.makedirs(os.path.dirname(path) or ".", exist_ok=True)
    doc.save(path)


def seed():
    # ── C-03: Block seed in production unless explicitly opted-in ──
    environment = os.getenv("ENVIRONMENT", "development").lower()
    allow_seed = os.getenv("ALLOW_DEMO_SEED", "false").lower() == "true"

    if environment == "production" and not allow_seed:
        print(
            "ERROR: Seed is blocked in production. "
            "Set ALLOW_DEMO_SEED=true to override.",
            file=sys.stderr,
        )
        sys.exit(1)

    # Get passwords from env (fall back to defaults only in development)
    admin_password = os.getenv(
        "SEED_ADMIN_PASSWORD",
        "admin123" if environment != "production" else None,
    )
    staff_password = os.getenv(
        "SEED_STAFF_PASSWORD",
        "staff123" if environment != "production" else None,
    )

    if not admin_password or not staff_password:
        print(
            "ERROR: Set SEED_ADMIN_PASSWORD and SEED_STAFF_PASSWORD env vars.",
            file=sys.stderr,
        )
        sys.exit(1)

    Base.metadata.create_all(bind=engine)
    os.makedirs(TEMPLATE_DIR, exist_ok=True)

    db = SessionLocal()
    try:
        ensure_appointment_letter(db)

        if db.query(models.User).count() > 0:
            print("Database already seeded. Skipping.")
            return

        admin = models.User(
            username="admin",
            full_name="Administrator",
            password_hash=hash_password(admin_password),
            role="admin",
        )
        staff = models.User(
            username="staff",
            full_name="Staff User",
            password_hash=hash_password(staff_password),
            role="staff",
        )
        db.add_all([admin, staff])

        doc_types = {}
        for name, slug in DOCUMENT_TYPES:
            dt = models.DocumentType(name=name, slug=slug)
            db.add(dt)
            db.flush()
            doc_types[slug] = dt

        country_map = {}
        trade_map = {}
        company_map = {}

        for code, data in SEED_DATA.items():
            country = models.Country(name=data["name"], code=code)
            db.add(country)
            db.flush()
            country_map[code] = country

            for trade_name, companies in data["trades"].items():
                trade = models.Trade(name=trade_name, country_id=country.id)
                db.add(trade)
                db.flush()
                trade_map[(code, trade_name)] = trade

                for company_name in companies:
                    company = models.Company(
                        name=company_name,
                        trade_id=trade.id,
                        country_id=country.id,
                    )
                    db.add(company)
                    db.flush()
                    company_map[(code, trade_name, company_name)] = company

        db.commit()

        sample_filename = "sample_canada_construction_buildrightcorp_offer_letter.docx"
        sample_path = os.path.join(TEMPLATE_DIR, sample_filename)
        create_sample_docx(sample_path)

        canada = country_map["ca"]
        construction = trade_map[("ca", "Construction Worker")]
        buildright = company_map[("ca", "Construction Worker", "BuildRight Corp")]
        offer_letter = doc_types["offer_letter"]

        template = models.Template(
            document_type_id=offer_letter.id,
            company_id=buildright.id,
            trade_id=construction.id,
            country_id=canada.id,
            docx_filename=sample_filename,
            label_overrides_json=json.dumps(LABEL_OVERRIDES),
            category="Construction and Infrastructure",
            format_slug="format1",
            format_label="Format 1 — Corporate Formal (Navy & Gold)",
            version=1,
            is_active=True,
        )
        db.add(template)

        contract_filename = "sample_canada_construction_buildrightcorp_employment_contract.docx"
        contract_path = os.path.join(TEMPLATE_DIR, contract_filename)
        create_employment_contract_docx(contract_path)
        employment_contract = doc_types["employment_contract"]
        contract_template = models.Template(
            document_type_id=employment_contract.id,
            company_id=buildright.id,
            trade_id=construction.id,
            country_id=canada.id,
            docx_filename=contract_filename,
            label_overrides_json=json.dumps(EMPLOYMENT_CONTRACT_LABELS),
            category="Construction and Infrastructure",
            format_slug="format1",
            format_label="Format 1 — Corporate Formal (Navy & Gold)",
            version=1,
            is_active=True,
        )
        db.add(contract_template)

        appt_filename = "sample_canada_construction_buildrightcorp_appointment_letter.docx"
        appt_path = os.path.join(TEMPLATE_DIR, appt_filename)
        create_appointment_letter_docx(appt_path)
        appointment_letter = doc_types["appointment_letter"]
        appt_template = models.Template(
            document_type_id=appointment_letter.id,
            company_id=buildright.id,
            trade_id=construction.id,
            country_id=canada.id,
            docx_filename=appt_filename,
            label_overrides_json=json.dumps(APPOINTMENT_LETTER_LABELS),
            category="Construction and Infrastructure",
            format_slug="format1",
            format_label="Format 1 — Corporate Formal (Navy & Gold)",
            version=1,
            is_active=True,
        )
        db.add(appt_template)

        nz = country_map.get("nz")
        if nz:
            nz_construction = trade_map.get(("nz", "Construction Worker"))
            kiwibuild = company_map.get(("nz", "Construction Worker", "KiwiBuild Ltd"))
            if nz_construction and kiwibuild:
                nz_filename = "sample_nz_construction_kiwibuild_offer_letter.docx"
                nz_path = os.path.join(TEMPLATE_DIR, nz_filename)
                if not os.path.exists(nz_path):
                    create_sample_docx(nz_path)
                nz_template = models.Template(
                    document_type_id=offer_letter.id,
                    company_id=kiwibuild.id,
                    trade_id=nz_construction.id,
                    country_id=nz.id,
                    docx_filename=nz_filename,
                    label_overrides_json=json.dumps(LABEL_OVERRIDES),
                    category="Construction and Infrastructure",
                    format_slug="format1",
                    format_label="Format 1 — Corporate Formal (Navy & Gold)",
                    version=1,
                    is_active=True,
                )
                db.add(nz_template)
        db.commit()

        print("Seed complete")
    finally:
        db.close()


if __name__ == "__main__":
    seed()

"""Trade-linked position pairing + trade search (name/synonym)."""
from __future__ import annotations

import io

from docx import Document

from services.trade_linked_position import DUTIES_REQUIRES_POSITION_MSG
from tests.test_phase3_platform import _setup_published_flow_with_field


def _upload_template(client, dt_id: int, name: str = "t.docx") -> int:
    doc = Document()
    doc.add_paragraph("{{duties_block}}")
    doc.add_paragraph("{{position}}")
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    up = client.post(
        f"/api/platform/{dt_id}/templates",
        files={
            "file": (
                name,
                buf.getvalue(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert up.status_code in (200, 201), up.text
    return up.json()["id"]


def test_list_trades_q_matches_name_and_synonym(dual_org_clients):
    client_a = dual_org_clients["client_a"]
    industry = client_a.post(
        "/api/platform/trade-industries",
        json={"name": "Electrical"},
    )
    assert industry.status_code == 201, industry.text
    ind_id = industry.json()["id"]

    created = client_a.post(
        "/api/platform/trades",
        json={
            "name": "Electrician",
            "duties_text": "Install wiring\nTest circuits",
            "industry_id": ind_id,
            "synonyms": ["Sparky", "Electrical fitter"],
        },
    )
    assert created.status_code == 201, created.text

    by_name = client_a.get("/api/platform/trades", params={"q": "Electrician"})
    assert by_name.status_code == 200
    assert any(t["name"] == "Electrician" for t in by_name.json())

    by_syn = client_a.get("/api/platform/trades", params={"q": "Sparky"})
    assert by_syn.status_code == 200
    rows = by_syn.json()
    assert len(rows) >= 1
    assert rows[0]["name"] == "Electrician"
    assert "Sparky" in (rows[0].get("synonyms") or [])

    miss = client_a.get("/api/platform/trades", params={"q": "PlumberXYZ"})
    assert miss.status_code == 200
    assert miss.json() == []


def test_trade_linked_position_stores_duties_field_key(dual_org_clients):
    client_a = dual_org_clients["client_a"]
    setup = _setup_published_flow_with_field(client_a, slug="tlp-store")
    step_id = setup["step_id"]

    duties = client_a.post(
        f"/api/platform/steps/{step_id}/fields",
        json={
            "field_key": "duties_block",
            "field_label": "Duties Block",
            "field_type": "text",
            "is_required": False,
            "is_auto_generated": False,
            "auto_config_json": None,
        },
    )
    assert duties.status_code == 201, duties.text

    position = client_a.post(
        f"/api/platform/steps/{step_id}/fields",
        json={
            "field_key": "position",
            "field_label": "Position",
            "field_type": "text",
            "is_required": True,
            "is_auto_generated": False,
            "auto_config_json": {
                "kind": "trade_linked_position",
                "duties_field_key": "duties_block",
            },
        },
    )
    assert position.status_code == 201, position.text
    body = position.json()
    assert body["auto_config_json"]["kind"] == "trade_linked_position"
    assert body["auto_config_json"]["duties_field_key"] == "duties_block"
    assert body["is_auto_generated"] is False


def test_publish_blocked_when_duties_block_without_trade_linked_position(
    dual_org_clients,
):
    client_a = dual_org_clients["client_a"]
    setup = _setup_published_flow_with_field(client_a, slug="tlp-pub-block")
    step_id = setup["step_id"]
    dt_id = setup["dt_id"]

    duties = client_a.post(
        f"/api/platform/steps/{step_id}/fields",
        json={
            "field_key": "duties_block",
            "field_label": "Duties Block",
            "field_type": "text",
            "is_required": False,
        },
    )
    assert duties.status_code == 201, duties.text

    draft = client_a.post(f"/api/platform/{dt_id}/flow/new-draft")
    assert draft.status_code == 201, draft.text
    draft_id = draft.json()["id"]

    pub = client_a.post(f"/api/platform/{draft_id}/publish")
    assert pub.status_code == 400, pub.text
    assert DUTIES_REQUIRES_POSITION_MSG in str(pub.json().get("detail", ""))


def test_publish_ok_with_trade_linked_position_and_duties(dual_org_clients):
    client_a = dual_org_clients["client_a"]
    setup = _setup_published_flow_with_field(client_a, slug="tlp-pub-ok")
    step_id = setup["step_id"]
    dt_id = setup["dt_id"]

    assert (
        client_a.post(
            f"/api/platform/steps/{step_id}/fields",
            json={
                "field_key": "duties_block",
                "field_label": "Duties",
                "field_type": "text",
            },
        ).status_code
        == 201
    )
    assert (
        client_a.post(
            f"/api/platform/steps/{step_id}/fields",
            json={
                "field_key": "position",
                "field_label": "Position",
                "field_type": "text",
                "auto_config_json": {
                    "kind": "trade_linked_position",
                    "duties_field_key": "duties_block",
                },
            },
        ).status_code
        == 201
    )

    draft = client_a.post(f"/api/platform/{dt_id}/flow/new-draft")
    assert draft.status_code == 201, draft.text
    pub = client_a.post(f"/api/platform/{draft.json()['id']}/publish")
    assert pub.status_code == 200, pub.text


def test_mapping_duties_block_rejected_without_trade_linked_position(
    dual_org_clients,
):
    client_a = dual_org_clients["client_a"]
    setup = _setup_published_flow_with_field(client_a, slug="tlp-map-block")
    step_id = setup["step_id"]
    dt_id = setup["dt_id"]

    client_a.post(
        f"/api/platform/steps/{step_id}/fields",
        json={
            "field_key": "duties_block",
            "field_label": "Duties",
            "field_type": "text",
        },
    )

    template_id = _upload_template(client_a, dt_id, "map-duties.docx")

    mapped = client_a.post(
        f"/api/platform/templates/{template_id}/mappings",
        json={
            "mappings": [
                {"placeholder_key": "duties_block", "field_key": "duties_block"},
            ]
        },
    )
    assert mapped.status_code == 400, mapped.text
    assert DUTIES_REQUIRES_POSITION_MSG in str(mapped.json().get("detail", ""))


def test_legacy_trade_linked_duties_kind_coerced_to_position(dual_org_clients):
    client_a = dual_org_clients["client_a"]
    setup = _setup_published_flow_with_field(client_a, slug="tlp-legacy")
    step_id = setup["step_id"]
    created = client_a.post(
        f"/api/platform/steps/{step_id}/fields",
        json={
            "field_key": "position",
            "field_label": "Position",
            "field_type": "text",
            "auto_config_json": {"kind": "trade_linked_duties"},
        },
    )
    assert created.status_code == 201, created.text
    cfg = created.json()["auto_config_json"]
    assert cfg["kind"] == "trade_linked_position"
    assert cfg["duties_field_key"] == "duties_block"

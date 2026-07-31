"""Org-scoped auto reference numbers + barcode at generate time."""
from __future__ import annotations

import io
import threading
import zipfile
from concurrent.futures import ThreadPoolExecutor
from datetime import datetime, timezone

from docx import Document
from sqlalchemy.orm import sessionmaker

import models
from models import OrgRefCounter
from services.org_ref_counter import format_org_ref_number, get_next_ref_number
from tests.test_phase3_platform import _setup_published_flow_with_field
from tests.test_template_delete import _map_complete
from tests.test_template_display_name import _upload


def test_format_org_ref_number_padding():
    assert format_org_ref_number("OLAW", 2026, 1) == "OLAW-2026-0001"
    assert format_org_ref_number("OLAW", 2026, 99) == "OLAW-2026-0099"
    assert format_org_ref_number("X", 2026, 10000) == "X-2026-10000"


def test_get_next_ref_number_sequential_and_year_reset(dual_org_clients):
    db = dual_org_clients["db"]
    org_id = dual_org_clients["org_a"]["org"].id
    client = dual_org_clients["client_a"]
    setup = _setup_published_flow_with_field(client, slug="ref-seq")
    dt_id = setup["dt_id"]

    a = get_next_ref_number(db, org_id, dt_id, "OLAW", year=2026)
    b = get_next_ref_number(db, org_id, dt_id, "OLAW", year=2026)
    c = get_next_ref_number(db, org_id, dt_id, "OLAW", year=2026)
    db.commit()
    assert a == "OLAW-2026-0001"
    assert b == "OLAW-2026-0002"
    assert c == "OLAW-2026-0003"

    d = get_next_ref_number(db, org_id, dt_id, "OLAW", year=2027)
    db.commit()
    assert d == "OLAW-2027-0001"


def test_get_next_ref_number_concurrent_unique(tmp_path):
    """
    Prove upsert returns unique sequences under concurrent writers.

    Uses a file SQLite DB (not the suite's StaticPool memory DB) so each
    thread can hold its own connection — closer to Postgres request isolation.
    """
    from sqlalchemy import create_engine, text

    db_path = tmp_path / "ref_race.db"
    race_engine = create_engine(
        f"sqlite:///{db_path}",
        connect_args={"check_same_thread": False, "timeout": 30},
    )
    with race_engine.begin() as conn:
        conn.execute(text("PRAGMA foreign_keys=OFF"))
        conn.execute(
            text(
                """
                CREATE TABLE org_ref_counters (
                    id INTEGER NOT NULL PRIMARY KEY AUTOINCREMENT,
                    org_id VARCHAR(36) NOT NULL,
                    document_type_id INTEGER NOT NULL,
                    year INTEGER NOT NULL,
                    last_sequence INTEGER NOT NULL DEFAULT 0,
                    CONSTRAINT uq_org_ref_counters_org_type_year
                        UNIQUE (org_id, document_type_id, year)
                )
                """
            )
        )

    SessionLocal = sessionmaker(bind=race_engine, autocommit=False, autoflush=False)
    org_id = "00000000-0000-0000-0000-0000000000aa"
    dt_id = 42
    results: list[str] = []
    lock = threading.Lock()
    errors: list[BaseException] = []

    def worker():
        db = SessionLocal()
        try:
            db.connection().exec_driver_sql("PRAGMA foreign_keys=OFF")
            value = get_next_ref_number(db, org_id, dt_id, "RACE", year=2099)
            db.commit()
            with lock:
                results.append(value)
        except BaseException as exc:  # noqa: BLE001
            db.rollback()
            with lock:
                errors.append(exc)
        finally:
            db.close()

    with ThreadPoolExecutor(max_workers=8) as pool:
        list(pool.map(lambda _: worker(), range(12)))

    assert not errors, errors
    assert len(results) == 12
    assert len(set(results)) == 12
    sequences = sorted(int(r.split("-")[-1]) for r in results)
    assert sequences == list(range(1, 13))


def test_cross_org_counters_independent(dual_org_clients):
    db = dual_org_clients["db"]
    client_a = dual_org_clients["client_a"]
    client_b = dual_org_clients["client_b"]
    setup_a = _setup_published_flow_with_field(client_a, slug="ref-orga")
    setup_b = _setup_published_flow_with_field(client_b, slug="ref-orgb")
    org_a = dual_org_clients["org_a"]["org"].id
    org_b = dual_org_clients["org_b"]["org"].id

    for _ in range(3):
        get_next_ref_number(db, org_a, setup_a["dt_id"], "AAA", year=2026)
    db.commit()

    first_b = get_next_ref_number(db, org_b, setup_b["dt_id"], "BBB", year=2026)
    db.commit()
    assert first_b == "BBB-2026-0001"

    row_a = (
        db.query(OrgRefCounter)
        .filter(
            OrgRefCounter.org_id == org_a,
            OrgRefCounter.document_type_id == setup_a["dt_id"],
            OrgRefCounter.year == 2026,
        )
        .first()
    )
    assert row_a is not None
    assert row_a.last_sequence == 3


def _docx_with_ref_and_barcode() -> bytes:
    doc = Document()
    doc.add_paragraph("Ref: {{ref_number}}")
    doc.add_paragraph("{{ref_number_barcode}}")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_generate_injects_ref_number_and_barcode(dual_org_clients):
    client = dual_org_clients["client_a"]
    setup = _setup_published_flow_with_field(client, slug="ref-gen", field_key="cand_name")
    dt_id = setup["dt_id"]
    step_id = setup["step_id"]

    # Add auto-ref field on the same (published) step
    auto = client.post(
        f"/api/platform/steps/{step_id}/fields",
        json={
            "field_key": "ref_number",
            "field_label": "Reference number",
            "field_type": "text",
            "is_required": True,  # server must force False for auto
            "is_auto_generated": True,
            "auto_config_json": {"kind": "ref_number", "prefix": "OLAW"},
        },
    )
    assert auto.status_code == 201, auto.text
    assert auto.json()["is_auto_generated"] is True
    assert auto.json()["is_required"] is False
    assert auto.json()["auto_config_json"]["prefix"] == "OLAW"

    files = {
        "file": (
            "ref.docx",
            _docx_with_ref_and_barcode(),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    }
    up = client.post(f"/api/platform/{dt_id}/templates", files=files)
    assert up.status_code == 201, up.text
    tmpl_id = up.json()["id"]

    mapped = client.post(
        f"/api/platform/templates/{tmpl_id}/mappings",
        json={
            "mappings": [
                {"placeholder_key": "ref_number", "field_key": "ref_number"},
                {"placeholder_key": "ref_number_barcode", "field_key": "ref_number"},
            ]
        },
    )
    assert mapped.status_code == 200, mapped.text

    year = datetime.now(timezone.utc).year
    gen = client.post(
        f"/api/platform/{dt_id}/generate",
        json={"template_id": tmpl_id, "fields": {setup["field_key"]: "Ada"}},
    )
    assert gen.status_code == 201, gen.text
    doc_id = gen.json()["document_id"]

    dl = client.get(f"/api/platform/generated/{doc_id}/download")
    assert dl.status_code == 200, dl.text
    expected = f"OLAW-{year}-0001"
    with zipfile.ZipFile(io.BytesIO(dl.content)) as zf:
        xml = zf.read("word/document.xml").decode("utf-8", errors="replace")
    assert expected in xml
    assert "a:blip" in xml or "w:drawing" in xml
    assert "{{ref_number_barcode}}" not in xml
    assert "{{ref_number}}" not in xml

    # Client-supplied ref_number must be ignored/overridden
    gen2 = client.post(
        f"/api/platform/{dt_id}/generate",
        json={
            "template_id": tmpl_id,
            "fields": {
                setup["field_key"]: "Bob",
                "ref_number": "HACKED-9999-9999",
            },
        },
    )
    assert gen2.status_code == 201, gen2.text
    dl2 = client.get(
        f"/api/platform/generated/{gen2.json()['document_id']}/download"
    )
    assert dl2.status_code == 200
    with zipfile.ZipFile(io.BytesIO(dl2.content)) as zf:
        xml2 = zf.read("word/document.xml").decode("utf-8", errors="replace")
    assert f"OLAW-{year}-0002" in xml2
    assert "HACKED" not in xml2


def test_wizard_visible_fields_contract():
    """Mirror frontend wizardVisibleFields — auto + barcode keys never render."""
    fields = [
        {"field_key": "cand_name", "is_auto_generated": False},
        {"field_key": "ref_number", "is_auto_generated": True},
        {"field_key": "ref_number_barcode", "is_auto_generated": False},
    ]
    visible = [
        f
        for f in fields
        if not f.get("is_auto_generated")
        and str(f.get("field_key") or "").lower() != "ref_number_barcode"
    ]
    assert [f["field_key"] for f in visible] == ["cand_name"]


def test_cannot_create_ref_number_barcode_as_separate_field(dual_org_clients):
    """Barcode must never be its own FieldDefinition — one auto-ref field only."""
    client = dual_org_clients["client_a"]
    setup = _setup_published_flow_with_field(client, slug="no-barcode-fd")
    step_id = setup["step_id"]

    blocked = client.post(
        f"/api/platform/steps/{step_id}/fields",
        json={
            "field_key": "ref_number_barcode",
            "field_label": "Ref Number Barcode",
            "field_type": "text",
            "is_required": True,
        },
    )
    assert blocked.status_code == 400, blocked.text
    assert "not a separate field" in blocked.json()["detail"].lower()

    # Creating the single auto-ref field still works
    ok = client.post(
        f"/api/platform/steps/{step_id}/fields",
        json={
            "field_key": "ref_number",
            "field_label": "Reference Number",
            "field_type": "text",
            "is_auto_generated": True,
            "auto_config_json": {"kind": "ref_number", "prefix": "OLAW"},
        },
    )
    assert ok.status_code == 201, ok.text
    assert ok.json()["field_key"] == "ref_number"
    assert ok.json()["is_auto_generated"] is True


def test_auto_field_excluded_from_required_validation(dual_org_clients):
    client = dual_org_clients["client_a"]
    setup = _setup_published_flow_with_field(client, slug="ref-req", field_key="cand_name")
    dt_id = setup["dt_id"]
    step_id = setup["step_id"]

    auto = client.post(
        f"/api/platform/steps/{step_id}/fields",
        json={
            "field_key": "ref_number",
            "field_label": "Ref",
            "field_type": "text",
            "is_auto_generated": True,
            "auto_config_json": {"prefix": "ZZ"},
        },
    )
    assert auto.status_code == 201

    up = _upload(client, dt_id, filename="plain.docx", placeholder="cand_name")
    tmpl_id = up.json()["id"]
    _map_complete(client, tmpl_id, setup["field_key"], placeholder="cand_name")

    gen = client.post(
        f"/api/platform/{dt_id}/generate",
        json={"template_id": tmpl_id, "fields": {setup["field_key"]: "Zed"}},
    )
    assert gen.status_code == 201, gen.text
